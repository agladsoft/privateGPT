import os
import docx
from datetime import datetime
from private_gpt.constants import PROJECT_ROOT_PATH


def build_form(message):
    pass


def create_doc(message, surname, name, post, department):
    doc = docx.Document()
    doc.add_paragraph("Генеральному директору").alignment = 2
    doc.add_paragraph("ООО «Рускон»").alignment = 2
    doc.add_paragraph("Колюху Н.И.\n\n").alignment = 2
    doc.add_paragraph(f"От {surname}").alignment = 2
    doc.add_paragraph(f"{name}").alignment = 2
    doc.add_paragraph(f"{post}").alignment = 2
    doc.add_paragraph(f"{department}\n\n").alignment = 2
    doc.add_paragraph('Заявление').alignment = 1
    doc.add_paragraph(f"{message}\n\n\n").alignment = 1
    doc.add_paragraph("_________").alignment = 2
    doc.add_paragraph("Титов С.С.").alignment = 2
    doc.add_paragraph(f"{datetime.now().strftime('%d.%m.%Y')}").alignment = 2
    doc.add_paragraph("«Согласовано»").alignment = 0
    doc.add_paragraph("Непосредственный руководитель").alignment = 0
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Директор Департамента стратегического и организационного развития"
    table.cell(0, 1).text = "В.Д. Шальнов"
    table.cell(0, 2).text = f"{datetime.now().strftime('%d.%m.%Y')}"
    path_dir = os.path.join(f"{PROJECT_ROOT_PATH}/upload_files", "files")
    os.makedirs(path_dir, exist_ok=True)
    path_file = f"{path_dir}/Заявка на отпуск.docx"
    doc.save(path_file)
    os.chmod(path_dir, 0o0777)
    return f'<a href="file/{path_file}" target="_blank" rel="noopener noreferrer">{os.path.basename(path_file)}</a>'
